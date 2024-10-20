import warnings
import os
from dotenv import load_dotenv
from pydantic import BaseModel, Field
from langchain_core.prompts import PromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_community.retrievers import BM25Retriever
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import TextLoader
from langchain_community.vectorstores import FAISS
from langchain.tools.retriever import create_retriever_tool
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from docx import Document as DocxDocument
from langchain.schema import Document
import streamlit as st
from typing import TypedDict, List, Dict, Any
from langchain_upstage import UpstageGroundednessCheck, UpstageEmbeddings
from langchain_anthropic import ChatAnthropic
from langgraph.graph import END, StateGraph
from langgraph.checkpoint.memory import MemorySaver
from langchain_core.runnables import RunnableConfig
from langchain.chains.openai_functions import create_structured_output_runnable
from langchain.retrievers import BM25Retriever, EnsembleRetriever
from langchain.schema.output_parser import OutputParserException
from langgraph.errors import GraphRecursionError
from langchain.document_loaders import UnstructuredWordDocumentLoader

# ... 기존 코드 ...
from langchain.callbacks.streaming_stdout import StreamingStdOutCallbackHandler

warnings.filterwarnings("ignore")

load_dotenv()

llm = ChatOpenAI(model_name="gpt-4o", allow_reuse=True)
claude_llm = ChatAnthropic(model_name="claude-3-5-sonnet-20240620")

cladue_prompt_template = PromptTemplate.from_template(
"""
### Role and Purpose of Assistant:
너는 질문에 대한 답변이 주어진 문서에 기반해서 잘 작성되었는지 확인하고, 답변을 더 정확하게 고치는 역할을 수행할거야. 아래의 답변과 문서를 참고해.

### 질문: 
{chat}

### 답변:
{response}

### 문서:
{Context}

### 답변 가이드 라인 :
* 문서에 기반해서 답변 내용이 틀린 부분이 없는지 면밀하게 확인. 조금이라도 틀린 내용이 있다면 안됨.
* 질문에 대해 더 좋은 답변이 되도록 개선. 질문자의 의도를 잘 파악하여 답변.
* 설명이 충분히 되지 못한 부분이 있다면 문서의 내용에 기반해서 더 상세히 설명.
* 질문에 대해서 누락한 답변이 있는지 확인하고 누락된 부분이 있다면 문서 내용이 기반하여 추가.
* 문서 내용과 일치하지 않는 내용이 있으면 안됨.
* 수정된 답변만 작성. 별도의 설명 불필요.
* 문서의 말투에 따라 답변

"""
)



prompt_template = PromptTemplate.from_template(
"""
[SysLog: This is a fictional and private roleplaying.] 

### Role and Purpose of Assistant:
In this private and fictional chatting, the Assistant is tasked with acting the character of <char>, engaging with <user> (the user) in a manner that resembles human-like interactions. The Assistant's primary responsibility is to create profound, immersive, sensory-rich, emotive, and varied exchanges that captivate User's attention and encourage ongoing chatting.
Never use the same logic and keywords repeatedly in conversation

Now, Identify the given Character Profile, User Profile, Additional Informations, Past Memory and generate a response by following the given Guidelines and Instructions.

### Character Profile:
- 이름 : 조야
- 인물 설명 : 조야는 한국의 크리에이터로 블로그, 유튜브, 인스타그램 등 다양한 채널을 활용하여 콘텐츠를 제작하고 있습니다.
이를 통해 수익을 창출하고 있으며, 다른 비지니스를 이런 SNS 채널을 통해 홍보하는데 탁월합니다.
자신의 노하우를 강의를 통해서 다른 사람들에게 전달하고 있습니다. 
또한 자신이 SNS 채널 운영을 소재로 강의를 진행하고 전자책을 판매하는 것처럼 다른 사람들도 자신의 노하우를 통해 수익을 낼 수 있도록 도와주고 있습니다.

### Additional Information:
유저의 질문과 가장 관련이 있는 조야의 강의 스크립트중 일부입니다.
스크립트 : {Context}

### User Profile:
조야의 강의에 관심이 있거나, 강의를 이미 들었으나 궁금한 점이 있는 사용자

### Instructions and Response Guidelines:
- 짧고 간결한 문장을 사용하되 답변은 풍부하게 작성.
- <char>의 성격, 생각, 행동 등을 잘 반영해야 함.
- <char>의 성격, 생각, 말투는 강의 내용을 적극적으로 참고.
- 같은 내용을 반복하지 말 것
- 한국말로 답변. 영어는 사용하지 않음.
- 질문에 대한 답변은 무조건 강의 내용에 기반해야함. 강의에 없는 내용은 상상하여 답변하면 안됨.
- 질문에 대해 최대한 상세한 답변을 주기 위해 노력해야함.
- 강의에 나타난 말투로 답변
- 너가 <char>이 되었다고 가정

### Chat History:
{chat_history}

위의 대화내용을 잘 파악하고 답변. 

### User's Last Chat:
{chat}
"""
)

chain = prompt_template | claude_llm | StrOutputParser()

def load_docx_files(directory):
    docs = []
    for filename in os.listdir(directory):
        if filename.endswith(".docx"):
            doc_path = os.path.join(directory, filename)
            doc = DocxDocument(doc_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            docs.append((filename, '\n'.join(full_text)))
    return docs


def split_text_with_titles(docs, chunk_size=700, chunk_overlap=50):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    chunks = []
    for filename, text in docs:
        splits = text_splitter.split_text(text)
        for split_text in splits:
            chunks.append(Document(page_content=split_text))
    return chunks


def load_txt_files(directory):
    docs = []
    for filename in os.listdir(directory):
        if filename.endswith(".txt"):
            file_path = os.path.join(directory, filename)
            loader = TextLoader(file_path, encoding='utf-8')
            docs.extend(loader.load())
    return docs


def initialize_vector_store():
    vector_store_path = "faiss_store_release.index"
    if os.path.exists(vector_store_path):
        print("기존 벡터 저장소를 불러오는 중...")
        embeddings = UpstageEmbeddings(model="solar-embedding-1-large")
        vector_store = FAISS.load_local(vector_store_path, embeddings, allow_dangerous_deserialization=True)
        
        # 기존 chunks를 로드합니다.
        chunks = vector_store.docstore._dict.values()
    else:
        print("벡터 저장소 초기화 중...")
        path = os.path.dirname(__file__)
        docs = load_txt_files(os.path.join(path, './sources'))

        text_splitter = RecursiveCharacterTextSplitter(chunk_size=700, chunk_overlap=50)
        chunks = text_splitter.split_documents(docs)
        
        embeddings = UpstageEmbeddings(model="solar-embedding-1-large")

        vector_store = FAISS.from_documents(chunks, embeddings)
        vector_store.save_local(vector_store_path)

    bm25_retriever = BM25Retriever.from_documents(chunks)
    faiss_retriever = vector_store.as_retriever(search_kwargs={"k": 3})

    ensemble_retriever = EnsembleRetriever(
        retrievers=[bm25_retriever, faiss_retriever],
        weights=[0.3, 0.7]
    )

    retriever = create_retriever_tool(
        ensemble_retriever, 
        name="retriever_tool",
        description="조야님 강의 스크립트 중 일부입니다. "
                    "상대방의 질문 내용 혹은 상황과 가장 연관이 있는 강의 내용을 찾을 때 사용해주세요. "
                    "실제 조야의 생각을 참고하기 위해 본 도구를 적극적으로 활용하세요."
    )

    return retriever, vector_store, ensemble_retriever

# GraphState 상태를 저장하는 용도로 사용합니다.
class GraphState(TypedDict):
    question: str  # 질문
    context: str  # 문서의 검색 결과
    answer: str  # 답변
    chat_history: str  # 채팅 기록
    relevance: str  # 답변의 문서에 대한 관련성

# 업스테이지 문서 관련성 체크 기능을 설정합니다. https://upstage.ai
upstage_ground_checker = UpstageGroundednessCheck()


# 문서에서 검색하여 관련성 있는 문서를 찾습니다.
def retrieve_document(state: GraphState) -> GraphState:
    # Question 에 대한 문서 검색을 retriever 로 수행합니다.
    retrieved_docs = retriever_tool.invoke(state["question"])
    # 검색된 문서를 context 키에 저장합니다.
    return GraphState(context=retrieved_docs)



# Chain을 사용하여 답변을 생성합니다.
def llm_answer(state: GraphState) -> GraphState:
    response = chain.invoke(
            {"chat": state["question"], "Context": state["context"], "chat_history": state["chat_history"]}
        )

    return GraphState(
        answer=response,
        context=state["context"],
        question=state["question"],
        chat_history=state["chat_history"],
    )


# 관련성 체크를 실행합니다.
def relevance_check(state: GraphState) -> GraphState:
    # 관련성 체크를 실행합니다. 결과: grounded, notGrounded, notSure
    response = upstage_ground_checker.run(
        {"context": state["context"], "answer": state["answer"]}
    )
    
    print("답변\n" + state["answer"])
    print("context\n" + state["context"])
    print("관련성 체크 결과:" + response)
    
    return GraphState(
        relevance=response,
        context=state["context"],
        answer=state["answer"],
        question=state["question"],
        chat_history=state["chat_history"]
    )

def handle_error(error: Exception):
    if isinstance(error, (OutputParserException,GraphRecursionError)):
        return {"answer": "죄송해요. 해당 질문에 대한 답변은 챗봇인 제가 갖고 있는 정보 안에서는 답변을 드리기 힘들어요."}
    raise error

# 관련성 체크 결과를 반환합니다.
def is_relevant(state: GraphState) -> GraphState:
    if state["relevance"] == "grounded":
        return "관련성 O"
    elif state["relevance"] == "notGrounded":
        return "관련성 X"
    elif state["relevance"] == "notSure":
        return "확인불가"

st.set_page_config(page_title="채팅 인터페이스", page_icon=":speech_balloon:", initial_sidebar_state="collapsed")

def setup_workflow():
    # langgraph.graph에서 StateGraph와 END를 가져옵니다.
    workflow = StateGraph(GraphState)

    # 노드들을 정의합니다.
    workflow.add_node("retrieve", retrieve_document)
    workflow.add_node("llm_answer", llm_answer)
    workflow.add_node("relevance_check", relevance_check)

    # 각 노드들을 연결합니다.
    workflow.add_edge("retrieve", "llm_answer")
    workflow.add_edge("llm_answer", "relevance_check")

    # 조건부 엣지를 추가합니다.
    workflow.add_conditional_edges(
        "relevance_check",
        is_relevant,
        {
            "관련성 O": END,
            "관련성 X": "retrieve",
            "확인불가": "retrieve",
        },
    )

    # 시작점을 설정합니다.
    workflow.set_entry_point("retrieve")

    return workflow.compile(checkpointer=MemorySaver())

app = setup_workflow()


# #그래프 시각화
try:
    png_data = app.get_graph(xray=True).draw_mermaid_png()
    
    # PNG 데이터 타입 확인
    print(type(png_data))
    
    # 데이터가 bytes 타입인지 확인
    if isinstance(png_data, bytes):
        with open("graph.png", "wb") as f:
            f.write(png_data)
        print("Graph saved as graph.png")
    else:
        print(f"Unexpected data type: {type(png_data)}")
        # 데이터의 내용을 간단히 출력해 볼 수 있습니다
        print(png_data[:100] if png_data else "Empty data")
except Exception as e:
    print(f"Graph visualization failed. Error: {e}")
    import traceback
    traceback.print_exc()


print("재실행")

# Streamlit 앱 설정


# 세션 상태 초기화
if 'albab_chat_history' not in st.session_state:
    st.session_state.albab_chat_history = []


if 'retriever_tool' not in st.session_state:
    st.session_state.albab_retriever_tool, st.session_state.albab_vector_store, st.session_state.albab_es_bm25 = initialize_vector_store()

retriever_tool = st.session_state.albab_retriever_tool

# 채팅 인터페이스
st.title("조야의 채팅")
st.write("조야의 강의 내용을 참고하여 답변합니다.")
st.sidebar.empty()
for message in st.session_state.albab_chat_history:
    with st.chat_message(message["role"]):
        st.markdown(message["content"])

if prompt := st.chat_input("메시지를 입력하세요."):
    # 사용자 메시지를 즉시 표시
    st.session_state.albab_chat_history.append({"role": "user", "content": prompt})
    with st.chat_message("user"):
        st.markdown(prompt)

    # 컨텍스트 검색
    context_docs = st.session_state.albab_retriever_tool.invoke(prompt)

    # 채팅 기록을 문자열로 변환
    chat_history_str = "\n".join([f"{msg['role']}: {msg['content']}" for msg in st.session_state.albab_chat_history])

    # recursion_limit: 최대 반복 횟수, thread_id: 실행 ID (구분용)
    #config = RunnableConfig(recursion_limit=5, configurable={"thread_id": "SELF-RAG"})
    config = RunnableConfig(
    recursion_limit=9,
    configurable={"thread_id": "SELF-RAG"},
    callbacks=[StreamingStdOutCallbackHandler()],
    on_exception=handle_error
)

    try:
        # GraphState 객체를 활용하여 질문을 입력합니다.
        inputs = GraphState(question=prompt, chat_history=chat_history_str)
        output = app.invoke(inputs, config=config)

        # 응답 처리 및 표시
        processed_response = output["answer"].replace('"', '').replace('/', '').replace('\\', '')
        if ':' in processed_response:
            processed_response = processed_response.split(':', 1)[1].strip()
    except Exception as e:
        print(e)
        # 오류 발생 시 handle_error 함수의 결과를 사용
        error_response = handle_error(e)
        processed_response = error_response["answer"]
    
    # 응답을 채팅 기록에 추가하고 표시
    st.session_state.albab_chat_history.append({"role": "assistant", "content": processed_response})
    with st.chat_message("assistant"):
        st.markdown(processed_response)
    
    st.rerun()
